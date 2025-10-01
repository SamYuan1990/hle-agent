import logging
import os

from docx import Document  # 添加docx库导入
from pdfminer.high_level import extract_text

# from pptx import Presentation


def is_pdf_file(filepath):
    """
    检查文件是否以.pdf结尾（不区分大小写）

    参数:
    filepath (str): 文件的绝对路径

    返回:
    bool: 如果是PDF文件返回True，否则返回False
    """
    # 使用os.path.splitext获取文件扩展名并转换为小写进行比较
    return os.path.splitext(filepath)[1].lower() == ".pdf"


def is_pptx_file(filepath):
    """
    检查文件是否以.pptx结尾（不区分大小写）

    参数:
    filepath (str): 文件的绝对路径

    返回:
    bool: 如果是PPTX文件返回True，否则返回False
    """
    # 使用os.path.splitext获取文件扩展名并转换为小写进行比较
    return os.path.splitext(filepath)[1].lower() == ".pptx"


def is_docx_file(filepath):
    """
    检查文件是否以.docx结尾（不区分大小写）

    参数:
    filepath (str): 文件的绝对路径

    返回:
    bool: 如果是DOCX文件返回True，否则返回False
    """
    # 使用os.path.splitext获取文件扩展名并转换为小写进行比较
    return os.path.splitext(filepath)[1].lower() == ".docx"


# def pptx_to_markdown_advanced(pptx_path):
#    prs = Presentation(pptx_path)
#    markdown_lines = []

#    for i, slide in enumerate(prs.slides):
#        markdown_lines.append(f"# Slide {i+1}\n\n")

#        if slide.shapes.title:
#            title = slide.shapes.title.text.strip()
#            markdown_lines.append(f"## {title}\n\n")

#        for shape in slide.shapes:
#            if hasattr(shape, "text") and shape != slide.shapes.title:
#                text = shape.text.strip()
#                if text:
#                    if any(char in text for char in ["•", "-", "*", "→"]):
#                        for line in text.split("\n"):
#                            if line.strip():
#                                cleaned = line.strip().lstrip("•-*→ ")
#                                markdown_lines.append(f"* {cleaned}\n")
#                        markdown_lines.append("\n")
#                    else:
#                        markdown_lines.append(f"{text}\n\n")

#            if shape.has_table:
#                table = shape.table
#                headers = [cell.text for cell in table.rows[0].cells]
#                markdown_lines.append(f"| {' | '.join(headers)} |\n")
#                markdown_lines.append(f"|{'|'.join(['---'] * len(headers))}|\n")
#                for row in table.rows[1:]:
#                    row_data = [cell.text for cell in row.cells]
#                    markdown_lines.append(f"| {' | '.join(row_data)} |\n")
#                markdown_lines.append("\n")

#        markdown_lines.append("---\n\n")

#    return "".join(markdown_lines)


def docx_to_text(docx_path):
    """
    将docx文件转换为纯文本

    参数:
    docx_path (str): docx文件的路径

    返回:
    str: 转换后的文本内容
    """
    try:
        doc = Document(docx_path)
        text_content = []

        # 处理段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 跳过空行
                text_content.append(paragraph.text)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:  # 只添加非空行
                    text_content.append(" | ".join(row_text))

        return "\n".join(text_content)

    except Exception as e:
        logging.error(f"处理DOCX文件时出错: {e}")
        return ""


def getfilecontent(filepath):
    file_content = ""
    try:
        ## pdf to text/markdown
        ## no wav to text/markdown as sst on client.(onnx is not cross plateform)
        if is_pdf_file(filepath):
            logging.info(f"{filepath} is a pdf file")
            file_content = extract_text(filepath)
            return file_content

        # if is_pptx_file(filepath):
        #    logging.info(f"{filepath} is a pptx file")
        #    file_content = pptx_to_markdown_advanced(filepath)
        #    return file_content

        if is_docx_file(filepath):
            logging.info(f"{filepath} is a docx file")
            file_content = docx_to_text(filepath)
            return file_content

        # default option
        logging.info(f"try {filepath} as text")
        with open(filepath, "r", encoding="utf-8") as f:
            file_content = f.read()
        return file_content

    except Exception as e:
        logging.error(f"处理文件 {filepath} 时出错: {e}")
        return file_content
